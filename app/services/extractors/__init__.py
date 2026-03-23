from __future__ import annotations

from pathlib import Path


def _load_docx_module():
    try:
        import docx
    except ModuleNotFoundError as exc:
        raise RuntimeError(
            "No se pudo cargar python-docx. Instala la dependencia `python-docx` para procesar archivos .docx."
        ) from exc
    return docx


def _load_openpyxl_loader():
    try:
        from openpyxl import load_workbook
    except ModuleNotFoundError as exc:
        raise RuntimeError(
            "No se pudo cargar openpyxl. Instala la dependencia `openpyxl` para procesar archivos .xlsx."
        ) from exc
    return load_workbook


def _load_ocr_stack():
    try:
        from PIL import Image
    except ModuleNotFoundError as exc:
        raise RuntimeError(
            "No se pudo cargar Pillow. Instala la dependencia `pillow` para procesar imágenes."
        ) from exc

    try:
        import pytesseract
    except ModuleNotFoundError as exc:
        raise RuntimeError(
            "No se pudo cargar pytesseract. Instala la dependencia `pytesseract` y el binario Tesseract OCR."
        ) from exc

    return Image, pytesseract


def _load_pymupdf():
    load_errors: list[str] = []

    try:
        import pymupdf

        if hasattr(pymupdf, "open"):
            return pymupdf
        load_errors.append("`pymupdf` se importó pero no expone `open`.")
    except Exception as exc:  # noqa: BLE001
        load_errors.append(f"pymupdf: {exc}")

    try:
        import fitz

        if hasattr(fitz, "open"):
            return fitz
        load_errors.append("`fitz` se importó pero no expone `open`; probablemente es el paquete incorrecto.")
    except Exception as exc:  # noqa: BLE001
        load_errors.append(f"fitz: {exc}")

    detail = " | ".join(load_errors) if load_errors else "sin detalle adicional"
    raise RuntimeError(
        "No se pudo cargar PyMuPDF para procesar PDFs. "
        "Si aparece un error como `from frontend import *`, tienes instalado el paquete `fitz` incorrecto. "
        "Solución recomendada: `pip uninstall -y fitz` y después `pip install --upgrade pymupdf`. "
        f"Detalle: {detail}"
    )


def _read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def _read_docx(path: Path) -> str:
    docx = _load_docx_module()
    document = docx.Document(str(path))
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    tables = []
    for table in document.tables:
        for row in table.rows:
            values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if values:
                tables.append(" | ".join(values))
    return "\n".join(paragraphs + tables)


def _read_xlsx(path: Path) -> str:
    load_workbook = _load_openpyxl_loader()
    workbook = load_workbook(str(path), data_only=True)
    parts: list[str] = []
    for sheet in workbook.worksheets:
        parts.append(sheet.title)
        for row in sheet.iter_rows(values_only=True):
            values = [str(value).strip() for value in row if value is not None and str(value).strip()]
            if values:
                parts.append(" | ".join(values))
    return "\n".join(parts)


def _read_pdf(path: Path) -> str:
    fitz = _load_pymupdf()
    with fitz.open(str(path)) as document:
        pages = [page.get_text("text") for page in document]
    return "\n".join(page.strip() for page in pages if page and page.strip())


def _read_image(path: Path) -> str:
    Image, pytesseract = _load_ocr_stack()
    image = Image.open(path)
    return pytesseract.image_to_string(image, lang="spa+eng").strip()


def _read_legacy_office(path: Path) -> str:
    raise ValueError(
        f"Formato {path.suffix.lower()} no soportado sin conversión externa (LibreOffice headless). "
        "Usa .docx o .xlsx para pruebas locales."
    )


EXTRACTORS = {
    ".txt": _read_text,
    ".docx": _read_docx,
    ".xlsx": _read_xlsx,
    ".pdf": _read_pdf,
    ".png": _read_image,
    ".jpg": _read_image,
    ".jpeg": _read_image,
    ".doc": _read_legacy_office,
    ".xls": _read_legacy_office,
    ".rtf": _read_legacy_office,
}


def extract_text_from_path(path: str | Path) -> tuple[str, dict[str, object]]:
    source = Path(path)
    extension = source.suffix.lower()
    extractor = EXTRACTORS.get(extension)
    if extractor is None:
        raise ValueError(f"Formato no soportado: {extension or 'sin extensión'}")
    text = extractor(source)
    return text, {"extension": extension, "extractor": extractor.__name__}
